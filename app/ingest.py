import os
import re
from urllib.parse import urlparse, parse_qs
import tempfile
from typing import Optional, Tuple

import httpx
from youtube_transcript_api import YouTubeTranscriptApi, NoTranscriptFound, TranscriptsDisabled
import yt_dlp

from .chunker import chunk_text
from .models import Item, ItemType, Chunk
from .storage import add_item, save_chunks, save_captions
from .llm import embed_texts, ensure_client
from .vector_store import add_chunks as vs_add_chunks
from .config import settings


YOUTUBE_REGEX = re.compile(r"(?:(?:[?&]v=)|(?:/embed/)|(?:/shorts/)|(?:youtu\.be/))([A-Za-z0-9_-]{11})")


def _extract_youtube_id(url: str) -> Optional[str]:
    # Robust parser that supports watch?v=, youtu.be/, /embed/, /shorts/
    try:
        u = urlparse(url)
        host = (u.netloc or "").lower()
        path = (u.path or "")
        # 1) Standard watch?v=
        qs = parse_qs(u.query or "")
        if "v" in qs and len(qs["v"]) > 0:
            vid = qs["v"][0]
            return vid[:11] if len(vid) >= 11 else vid
        # 2) youtu.be/<id>
        if "youtu.be" in host:
            parts = path.strip("/").split("/")
            if parts and len(parts[0]) >= 11:
                return parts[0][:11]
        # 3) /embed/<id>
        parts = path.strip("/").split("/")
        if len(parts) >= 2 and parts[0] in {"embed", "shorts", "live"}:
            return parts[1][:11]
    except Exception:
        pass
    # 4) Fallback regex
    m = YOUTUBE_REGEX.search(url)
    return m.group(1) if m else None


def _vtt_to_text(path: str) -> str:
    import re
    import html as html_lib

    out_lines = []
    tag_re = re.compile(r"<[^>]+>")  # remove <c> and <00:..> etc.
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            for line in f:
                s = line.strip()
                if not s:
                    continue
                if s.startswith("WEBVTT"):
                    continue
                if "-->" in s:
                    continue
                if s.isdigit():
                    continue
                if s.lower().startswith("kind:") or s.lower().startswith("language:"):
                    continue
                s = tag_re.sub(" ", s)
                s = html_lib.unescape(s)
                s = re.sub(r"\s+", " ", s).strip()
                if s:
                    out_lines.append(s)
    except Exception:
        return ""
    return " ".join(out_lines)


def _vtt_to_segments(path: str):
    import re
    import html as html_lib
    pattern = re.compile(r"(?P<start>\d{2}:\d{2}:\d{2}\.\d{3})\s+-->\s+(?P<end>\d{2}:\d{2}:\d{2}\.\d{3})")
    tag_re = re.compile(r"<[^>]+>")
    def to_seconds(hhmmss: str) -> float:
        hh, mm, ss = hhmmss.split(":")
        return int(hh) * 3600 + int(mm) * 60 + float(ss)
    segments = []
    start = None
    end = None
    buf: list[str] = []
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        for raw in f:
            line = raw.strip()
            if not line:
                continue
            if line.startswith("WEBVTT"):
                continue
            m = pattern.match(line)
            if m:
                if start is not None and buf:
                    text = tag_re.sub(" ", " ".join(buf))
                    text = html_lib.unescape(text)
                    text = re.sub(r"\s+", " ", text).strip()
                    if text:
                        segments.append({"start": start, "end": end if end is not None else start, "text": text})
                start = to_seconds(m.group("start"))
                end = to_seconds(m.group("end"))
                buf = []
                continue
            # skip index lines and metadata
            if line.isdigit() or line.lower().startswith(("kind:", "language:")):
                continue
            if "-->" in line:
                continue
            clean = tag_re.sub(" ", line)
            clean = html_lib.unescape(clean)
            clean = re.sub(r"\s+", " ", clean).strip()
            if clean:
                buf.append(clean)
    if start is not None and buf:
        text = tag_re.sub(" ", " ".join(buf))
        import re as _re
        text = html_lib.unescape(text)
        text = _re.sub(r"\s+", " ", text).strip()
        if text:
            segments.append({"start": start, "end": end if end is not None else start, "text": text})
    return segments


def _merge_segments(segments, max_chars: int = 600, max_gap_s: float = 2.5, max_span_s: float = 60.0):
    merged_texts = []
    metas = []
    if not segments:
        return merged_texts, metas
    cur_start = segments[0]["start"]
    cur_end = segments[0]["end"]
    cur_text = segments[0]["text"]
    for seg in segments[1:]:
        st, en, txt = seg["start"], seg["end"], seg["text"]
        span_ok = (en - cur_start) <= max_span_s
        gap_ok = (st - cur_end) <= max_gap_s
        chars_ok = (len(cur_text) + 1 + len(txt)) <= max_chars
        if span_ok and gap_ok and chars_ok:
            cur_end = en
            cur_text = f"{cur_text} {txt}"
        else:
            merged_texts.append(cur_text)
            metas.append({"start_s": float(cur_start), "end_s": float(cur_end)})
            cur_start, cur_end, cur_text = st, en, txt
    merged_texts.append(cur_text)
    metas.append({"start_s": float(cur_start), "end_s": float(cur_end)})
    return merged_texts, metas


def _try_ytdlp_autosubs(url: str) -> str:
    """Try to fetch auto-generated subtitles using yt-dlp without ffmpeg."""
    tmpdir = tempfile.mkdtemp()
    ydl_opts = {
        "skip_download": True,
        "writesubtitles": True,
        "writeautomaticsub": True,
        "subtitleslangs": ["en", "en.*"],
        "subtitlesformat": "vtt",
        "outtmpl": os.path.join(tmpdir, "%(id)s.%(ext)s"),
        # No ffmpeg needed for just subtitles
    }
    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            ydl.download([url])
        # Find any .vtt
        for f in os.listdir(tmpdir):
            if f.lower().endswith(".vtt"):
                path = os.path.join(tmpdir, f)
                txt = _vtt_to_text(path)
                if txt.strip():
                    return txt
    except Exception:
        return ""
    return ""

def _ytdlp_autosubs_segments(url: str):
    tmpdir = tempfile.mkdtemp()
    ydl_opts = {
        "skip_download": True,
        "writesubtitles": True,
        "writeautomaticsub": True,
        "subtitleslangs": ["en", "en.*"],
        "subtitlesformat": "vtt",
        "outtmpl": os.path.join(tmpdir, "%(id)s.%(ext)s"),
    }
    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            ydl.download([url])
        for f in os.listdir(tmpdir):
            if f.lower().endswith(".vtt"):
                path = os.path.join(tmpdir, f)
                return _vtt_to_segments(path)
    except Exception:
        return []
    return []


def ingest_youtube(board_id: str, url: str, title_hint: str = "") -> Item:
    vid = _extract_youtube_id(url)
    if not vid:
        raise ValueError("Invalid YouTube URL")

    # Try official captions first
    text = ""
    segments = []
    try:
        tracks = YouTubeTranscriptApi.list_transcripts(vid)
        transcript = tracks.find_transcript(["en"]) if tracks else None
        if transcript:
            entries = transcript.fetch()
            for e in entries:
                t = (e.get("text") or "").strip()
                if not t:
                    continue
                st = float(e.get("start", 0.0))
                en = st + float(e.get("duration", 0.0))
                segments.append({"start": st, "end": en, "text": t})
            text = " ".join([s["text"] for s in segments])
    except (NoTranscriptFound, TranscriptsDisabled, Exception):
        text = ""
        segments = []

    # Fallback 1: try yt-dlp auto subtitles
    if not text.strip():
        segments = _ytdlp_autosubs_segments(url)
        text = " ".join([s["text"] for s in segments])

    # Fallback 2: download audio and transcribe with Whisper
    if not text.strip():
        ensure_client()
        tmpdir = tempfile.mkdtemp()
        # Use ffmpeg-based postprocessing to extract m4a (requires ffmpeg in PATH)
        ydl_opts = {
            "format": "bestaudio/best",
            "outtmpl": os.path.join(tmpdir, "%(id)s.%(ext)s"),
            "postprocessors": [
                {
                    "key": "FFmpegExtractAudio",
                    "preferredcodec": "m4a",
                }
            ],
            # If ffmpeg is not in PATH, set "ffmpeg_location" here
        }
        outfile = None
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=True)
            # pick produced m4a in tmpdir
            for f in os.listdir(tmpdir):
                if f.lower().endswith(".m4a"):
                    outfile = os.path.join(tmpdir, f)
                    break

        if not outfile:
            raise RuntimeError("Failed to produce audio file for transcription (ffmpeg)")

        client = ensure_client()
        # Retry a couple of times in case of transient 500s
        last_err = None
        for _ in range(2):
            try:
                with open(outfile, "rb") as f:
                    tr = client.audio.transcriptions.create(model="whisper-1", file=f, response_format="verbose_json")
                    text = (getattr(tr, "text", "") or "").strip()
                    segs = getattr(tr, "segments", None) or []
                    segments = [{"start": float(s.get("start", 0.0)), "end": float(s.get("end", 0.0)), "text": (s.get("text") or "").strip()} for s in segs if (s.get("text") or "").strip()]
                    if text:
                        break
            except Exception as e:
                last_err = e
        if not text:
            # Fallback to gpt-4o-mini-transcribe if Whisper keeps failing
            try:
                with open(outfile, "rb") as f:
                    tr = client.audio.transcriptions.create(
                        model="gpt-4o-mini-transcribe",
                        file=f,
                        response_format="verbose_json",
                    )
                    text = (getattr(tr, "text", "") if not isinstance(tr, str) else tr).strip()
                    segs = getattr(tr, "segments", None) or []
                    if isinstance(segs, list) and segs:
                        segments = [{"start": float(s.get("start", 0.0)), "end": float(s.get("end", 0.0)), "text": (s.get("text") or "").strip()} for s in segs if (s.get("text") or "").strip()]
            except Exception:
                pass

        if not text:
            if last_err:
                raise last_err
            raise RuntimeError("Transcription failed (both Whisper and 4o-mini-transcribe)")

    title = title_hint or f"YouTube {vid}"
    item = Item(board_id=board_id, type=ItemType.YOUTUBE, title=title, source=url)
    item = add_item(item)

    # If we have timestamped segments, save them and create merged chunks with start/end
    if segments:
        save_captions(item.id, segments)
        texts, metas = _merge_segments(segments)
        if texts:
            embs = embed_texts(texts)
            meta_list = [{"item_id": item.id, **m} for m in metas]
            vs_add_chunks(item.id, texts, embs, meta_list)
            save_chunks([Chunk(item_id=item.id, text=t, order=i, start_s=metas[i]["start_s"], end_s=metas[i]["end_s"]) for i, t in enumerate(texts)])
        return item

    # Fallback: plain text (no timestamps)
    parts = chunk_text(text)
    if parts:
        embs = embed_texts(parts)
        vs_add_chunks(item.id, parts, embs, [{"item_id": item.id} for _ in parts])
        save_chunks([Chunk(item_id=item.id, text=p, order=i) for i, p in enumerate(parts)])
    return item


def ingest_text_document(board_id: str, title: str, text: str, source: str) -> Item:
    item = Item(board_id=board_id, type=ItemType.DOCUMENT, title=title, source=source)
    item = add_item(item)
    parts = chunk_text(text)
    if parts:
        embs = embed_texts(parts)
        vs_add_chunks(item.id, parts, embs)
        save_chunks([Chunk(item_id=item.id, text=p, order=i) for i, p in enumerate(parts)])
    return item


def ingest_web_url(board_id: str, url: str, title_hint: str = "") -> Item:
    from bs4 import BeautifulSoup
    from readability import Document

    resp = httpx.get(url, timeout=30.0)
    resp.raise_for_status()
    doc = Document(resp.text)
    title = title_hint or doc.short_title() or url
    html = doc.summary()
    soup = BeautifulSoup(html, "lxml")
    text = soup.get_text(" ")
    return ingest_text_document(board_id, title, text, url)


def ingest_pdf(board_id: str, file_path: str, title_hint: str = "") -> Item:
    import pypdf

    # First attempt with pypdf
    try:
        reader = pypdf.PdfReader(file_path)
        text = " \n".join(page.extract_text() or "" for page in reader.pages)
    except Exception:
        text = ""

    # Fallback to pdfminer.six if very little text was extracted
    if len((text or "").strip()) < 40:
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract_text  # type: ignore
            text2 = pdfminer_extract_text(file_path) or ""
            if len(text2.strip()) > len(text.strip()):
                text = text2
        except Exception:
            pass
    title = title_hint or os.path.basename(file_path)
    return ingest_text_document(board_id, title, text, file_path)


def ingest_docx(board_id: str, file_path: str, title_hint: str = "") -> Item:
    text = ""
    try:
        from docx import Document as Docx
        doc = Docx(file_path)
        parts = []
        # Paragraphs
        parts.extend(p.text for p in doc.paragraphs if p.text)
        # Tables
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        # Headers / Footers
        try:
            for section in doc.sections:
                if section.header and section.header.paragraphs:
                    parts.extend(p.text for p in section.header.paragraphs if p.text)
                if section.footer and section.footer.paragraphs:
                    parts.extend(p.text for p in section.footer.paragraphs if p.text)
        except Exception:
            pass
        text = "\n".join(parts)
    except Exception:
        text = ""
    # Fallback to docx2txt if available
    if len(text.strip()) < 10:
        try:
            import docx2txt  # type: ignore
            text2 = docx2txt.process(file_path) or ""
            if len(text2.strip()) > len(text.strip()):
                text = text2
        except Exception:
            pass
    # Fallback to mammoth if available (DOCX -> HTML -> text)
    if len(text.strip()) < 10:
        try:
            import mammoth  # type: ignore
            with open(file_path, "rb") as f:
                res = mammoth.convert_to_html(f)
                html = res.value or ""
            from bs4 import BeautifulSoup
            text2 = BeautifulSoup(html, "lxml").get_text(" ")
            if len(text2.strip()) > len(text.strip()):
                text = text2
        except Exception:
            pass
    # Fallback to raw XML unzip (no extra deps)
    if len(text.strip()) < 10:
        try:
            import zipfile
            from xml.etree import ElementTree as ET
            txt_parts = []
            with zipfile.ZipFile(file_path) as z:
                # Main document
                for name in z.namelist():
                    if name.startswith("word/") and (name.endswith("document.xml") or name.endswith("header1.xml") or name.endswith("header2.xml") or name.endswith("footer1.xml") or name.endswith("footer2.xml")):
                        try:
                            xml = z.read(name).decode("utf-8", errors="ignore")
                            root = ET.fromstring(xml)
                            # Extract all w:t nodes
                            for el in root.iter():
                                if el.tag.endswith('}t') and el.text:
                                    txt_parts.append(el.text)
                        except Exception:
                            continue
            if txt_parts:
                text2 = "\n".join(txt_parts)
                if len(text2.strip()) > len(text.strip()):
                    text = text2
        except Exception:
            pass
    title = title_hint or os.path.basename(file_path)
    return ingest_text_document(board_id, title, text, file_path)


def ingest_txt(board_id: str, file_path: str, title_hint: str = "") -> Item:
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
    except Exception:
        with open(file_path, "r", encoding="latin-1", errors="ignore") as f:
            text = f.read()
    title = title_hint or os.path.basename(file_path)
    return ingest_text_document(board_id, title, text, file_path)


def ingest_media(board_id: str, file_path: str, title_hint: str = "") -> Item:
    # Transcribe local audio/video
    ensure_client()
    client = ensure_client()
    segments = []
    text = ""
    with open(file_path, "rb") as f:
        tr = client.audio.transcriptions.create(model="whisper-1", file=f, response_format="verbose_json")
        text = getattr(tr, "text", "") or ""
        segs = getattr(tr, "segments", None) or []
        if isinstance(segs, list) and segs:
            segments = [{"start": float(s.get("start", 0.0)), "end": float(s.get("end", 0.0)), "text": (s.get("text") or "").strip()} for s in segs if (s.get("text") or "").strip()]
    title = title_hint or os.path.basename(file_path)
    item = Item(board_id=board_id, type=ItemType.AUDIOVIDEO, title=title, source=file_path)
    item = add_item(item)
    if segments:
        save_captions(item.id, segments)
        texts, metas = _merge_segments(segments)
        if texts:
            embs = embed_texts(texts)
            meta_list = [{"item_id": item.id, **m} for m in metas]
            vs_add_chunks(item.id, texts, embs, meta_list)
            save_chunks([Chunk(item_id=item.id, text=t, order=i, start_s=metas[i]["start_s"], end_s=metas[i]["end_s"]) for i, t in enumerate(texts)])
        return item
    # Fallback plain text
    parts = chunk_text(text)
    if parts:
        embs = embed_texts(parts)
        vs_add_chunks(item.id, parts, embs, [{"item_id": item.id} for _ in parts])
        save_chunks([Chunk(item_id=item.id, text=p, order=i) for i, p in enumerate(parts)])
    return item


